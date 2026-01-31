import gc
import os

import docx
import pdfplumber
import pypandoc
from odf.opendocument import load
from odf.text import P


class TextExtractor:
    def __init__(self, filepath):
        self.filepath = filepath
        self.extension = os.path.splitext(filepath)[-1].lower()

    def extract_file(self):
        """Default extraction method"""
        try:
            return pypandoc.convert_file(
                self.filepath, "plain", extra_args=["--wrap=none"]
            )
        except Exception as e:
            print(f"[WARN] Pandoc failed for {self.filepath} → {e}")
            return self.fallback_extract()

    def extract_file_large(self, chunk_size=1000000):  # ~1MB chunks for large files
        """Memory-efficient extraction for large files"""
        try:
            if self.extension == ".pdf":
                return self._extract_pdf_large()
            elif self.extension == ".txt":
                return self._extract_txt_large(chunk_size)
            elif self.extension in [".docx", ".doc"]:
                return self._extract_docx_large()
            else:
                # For other formats, use existing method but with memory cleanup
                result = self.extract_file()
                gc.collect()
                return result
        except Exception as e:
            print(f"Large file extraction failed: {e}")
            return self.fallback_extract()

    def fallback_extract(self):
        if self.extension == ".pdf":
            return self._extract_pdf()
        elif self.extension == ".docx":
            return self._extract_docx()
        elif self.extension == ".odt":
            return self._extract_odt()
        elif self.extension in [".txt", ".md", ".rtf"]:
            return self._extract_txt()
        else:
            raise ValueError(f"Unsupported file type: {self.extension}")

    def save(self, outdir="src/TTS/artifacts"):
        """Extracts text and saves it as .txt"""
        os.makedirs(outdir, exist_ok=True)
        base = os.path.splitext(os.path.basename(self.filepath))[0]
        outpath = os.path.join(outdir, f"{base}.txt")

        text = self.extract_file()
        with open(outpath, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"Saved {self.filepath} → {outpath}")
        return outpath

    def _extract_pdf(self):
        text = ""
        with pdfplumber.open(self.filepath) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text.strip()

    def _extract_pdf_large(self):
        """Memory-efficient PDF extraction with periodic cleanup"""
        text_parts = []
        try:
            with pdfplumber.open(self.filepath) as pdf:
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)

                    # Clear memory every 10 pages
                    if i % 10 == 0:
                        gc.collect()

            return "".join(text_parts)
        except Exception as e:
            print(f"PDF large extraction failed: {e}")
            return self._extract_pdf()

    def _extract_docx(self):
        doc = docx.Document(self.filepath)

        for section in doc.sections:
            header = section.header
            footer = section.footer

            # Clear all header paragraphs
            for paragraph in header.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

            # Clear all footer paragraphs
            for paragraph in footer.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)

        return "\n".join([p.text for p in doc.paragraphs])

    def _extract_docx_large(self):
        """Memory-efficient DOCX extraction"""
        try:
            # For very large DOCX files, process in chunks if possible
            return self._extract_docx()
        except Exception as e:
            print(f"DOCX large extraction failed: {e}")
            return self._extract_docx()

    def _extract_odt(self):
        doc = load(self.filepath)
        paragraphs = doc.getElementsByType(P)
        return "\n".join([p.firstChild.data for p in paragraphs if p.firstChild])

    def _extract_txt(self):
        with open(self.filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def _extract_txt_large(self, chunk_size):
        """Memory-efficient text file reading"""
        text_parts = []
        try:
            with open(self.filepath, "r", encoding="utf-8", errors="ignore") as f:
                while True:
                    chunk = f.read(chunk_size)
                    if not chunk:
                        break
                    text_parts.append(chunk)
            return "".join(text_parts)
        except Exception as e:
            print(f"Text large extraction failed: {e}")
            return self._extract_txt()


if __name__ == "__main__":
    files = [
        r"C:\Users\Hp\Downloads\Proof Of Concept (1).pdf",
        r"C:\Users\Hp\Downloads\AI-Powered Drawing Tool Research_.docx",
        r"src\TTS\Files\resa.md",
        r"src\TTS\Files\Sway.epub",
        r"src\TTS\Files\indes.html",
        r"src\TTS\Files\ram.txt",
    ]

    for f in files:
        try:
            extractor = TextExtractor(f)
            extractor.save()
        except Exception as e:
            print(f"Failed on {f}: {e}")
